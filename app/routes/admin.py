from __future__ import annotations
from app.utils.test_generator import choose_questions_exact

from io import BytesIO
from pathlib import Path
from uuid import uuid4
import random

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
import random

from flask import Blueprint, abort, current_app, flash, redirect, render_template, request, send_file, url_for
from flask_login import current_user, login_required
from openpyxl import Workbook
from werkzeug.utils import secure_filename
from flask_wtf import FlaskForm

from app.extensions import db
from app.forms import CreateTestForm, CreateUserForm, EditUserForm, QUESTION_TYPE_CHOICES, UploadTestsForm, UploadGroupForm
from app.models import ActionLog, AnswerOption, Attempt, Question, Test, TestAssignment, User, StudentGroup
from app.utils.i18n import translate as _t
from app.utils.parsers import ParseError, parse_uploaded_tests, parse_group_excel
from app.utils.seed import seed_demo_data


admin_bp = Blueprint('admin', __name__, url_prefix='/admin')


QUESTION_TYPE_LABEL_KEYS = {
    'single_choice': 'question_type.single_choice',
    'multiple_choice': 'question_type.multiple_choice',
    'matching': 'question_type.matching',
    'text': 'question_type.text',
    'ordering': 'question_type.ordering',
    'single': 'question_type.single_choice',
    'multiple': 'question_type.multiple_choice',
}

def build_attempt_risk_reasons(attempt: Attempt) -> list[str]:
    reasons: list[str] = []

    if (attempt.tab_switch_count or 0) >= 1:
        reasons.append(f'Переключал вкладки: {attempt.tab_switch_count or 0}')

    if (attempt.focus_loss_count or 0) >= 1:
        reasons.append(f'Терял фокус окна: {attempt.focus_loss_count or 0}')

    if (attempt.fullscreen_exit_count or 0) >= 1:
        reasons.append(f'Выходил из полноэкранного режима: {attempt.fullscreen_exit_count or 0}')

    if (attempt.heartbeat_count or 0) == 0:
        reasons.append('Нет heartbeat-сигналов')

    if (attempt.suspicion_score or 0) >= 1:
        reasons.append(f'Итоговый suspicion score: {attempt.suspicion_score or 0}')

    if not reasons:
        reasons.append('Подозрительных признаков не найдено')

    return reasons

def log_action(action_type: str, entity_type: str, description: str, entity_id: int | None = None) -> None:
    if not current_user.is_authenticated:
        return

    db.session.add(
        ActionLog(
            actor_user_id=current_user.id,
            action_type=action_type,
            entity_type=entity_type,
            entity_id=entity_id,
            description=description,
        )
    )


def staff_required():
    if not current_user.is_authenticated or not current_user.is_staff:
        abort(403)

def get_teacher_group_ids() -> list[int]:
    if current_user.is_admin:
        return [group.id for group in StudentGroup.query.all()]
    return [group.id for group in current_user.teaching_groups]


def get_teacher_student_ids() -> list[int]:
    if current_user.is_admin:
        return [user.id for user in User.query.filter_by(role='student').all()]

    group_ids = get_teacher_group_ids()
    if not group_ids:
        return []

    students = User.query.filter(
        User.role == 'student',
        User.group_id.in_(group_ids)
    ).all()
    return [student.id for student in students]


def get_teacher_test_ids() -> list[int]:
    if current_user.is_admin:
        return [test.id for test in Test.query.all()]

    tests = Test.query.filter_by(teacher_id=current_user.id).all()
    return [test.id for test in tests]

@admin_bp.route('/')
@login_required
def admin_dashboard():
    staff_required()

    if current_user.is_admin:
        stats = {
            'students': User.query.filter_by(role='student').count(),
            'tests': Test.query.count(),
            'assignments': TestAssignment.query.count(),
            'attempts': Attempt.query.filter(Attempt.finished_at.isnot(None)).count(),
        }
        tests_list = Test.query.order_by(Test.id.desc()).all()
    else:
        teacher_group_ids = get_teacher_group_ids()
        teacher_student_ids = get_teacher_student_ids()
        teacher_test_ids = get_teacher_test_ids()

        stats = {
            'students': len(teacher_student_ids),
            'tests': len(teacher_test_ids),
            'assignments': TestAssignment.query.filter(TestAssignment.test_id.in_(teacher_test_ids)).count() if teacher_test_ids else 0,
            'attempts': Attempt.query.filter(
                Attempt.finished_at.isnot(None),
                Attempt.test_id.in_(teacher_test_ids)
            ).count() if teacher_test_ids else 0,
        }

        tests_list = Test.query.filter(Test.id.in_(teacher_test_ids)).order_by(Test.id.desc()).all() if teacher_test_ids else []

    return render_template('admin/dashboard.html', stats=stats, tests=tests_list)


@admin_bp.route('/upload', methods=['GET', 'POST'])
@login_required
def upload_tests():
    staff_required()
    form = UploadTestsForm()

    if request.method == 'POST' and not form.validate_on_submit():
        for field_name, errors in form.errors.items():
            for error in errors:
                flash(f'{field_name}: {error}', 'danger')
        return render_template('admin/upload_tests.html', form=form)

    if form.validate_on_submit():
        file = form.file.data
        original_name = secure_filename(file.filename or 'tests.txt')
        extension = Path(original_name).suffix.lower()
        unique_name = f'{uuid4().hex}{extension}'

        upload_dir = Path(current_app.instance_path) / current_app.config['UPLOAD_FOLDER']
        upload_dir.mkdir(parents=True, exist_ok=True)

        saved_path = upload_dir / unique_name
        file.save(saved_path)

        try:
            parsed_tests = parse_uploaded_tests(str(saved_path))

            created_tests, created_test_ids = create_tests_from_parsed_data(
                parsed_tests,
                time_limit_minutes=form.time_limit_minutes.data,
                max_attempts=form.max_attempts.data,
                created_by=current_user.id,
                teacher_id=current_user.id,
            )

        except ParseError as error:
            saved_path.unlink(missing_ok=True)
            flash(_t('admin.parse_error', error=error), 'danger')
            return render_template('admin/upload_tests.html', form=form)

        log_action(
            action_type='import_tests',
            entity_type='test_import',
            entity_id=None,
            description=f'Импортировал тесты из файла "{original_name}", создано тестов: {created_tests}',
        )
        db.session.commit()

        flash(_t('admin.upload_success', filename=original_name, count=created_tests), 'success')
        return redirect(url_for('admin.admin_dashboard'))

    return render_template('admin/upload_tests.html', form=form)


@admin_bp.route('/upload/template')
@login_required
def download_test_template():
    staff_required()

    template_text = """Test: Matematika namunasi

2 ballik savol
Savol: 2+2 nechiga teng?
Turi: bir toʻgʻri javobli
Javoblar: 3 | *4 | 5 | 6

1 ballik savol
Savol: O‘zbekiston poytaxti?
Turi: matn
Javob: *Toshkent

Test: Informatika namunasi

1 ballik savol
Savol: HTML nimani bildiradi?
Turi: bir toʻgʻri javobli
Javoblar: Hyper Tool Markup Language | *HyperText Markup Language | HighText Markdown Language | Home Tool Markup Language

1 ballik savol
Savol: Python nima?
Turi: matn
Javob: *Dasturlash tili
"""

    buffer = BytesIO(template_text.encode('utf-8'))

    return send_file(
        buffer,
        as_attachment=True,
        download_name=_t('admin.template_filename'),
        mimetype='text/plain; charset=utf-8',
    )
@admin_bp.route('/users')
@login_required
def users():
    admin_required()

    if current_user.is_admin:
        teachers = User.query.filter_by(role='teacher').order_by(User.full_name.asc()).all()
        ungrouped_students = User.query.filter_by(role='student', group_id=None).order_by(User.username.asc()).all()
        groups = StudentGroup.query.order_by(StudentGroup.name.asc()).all()
        all_users = User.query.order_by(User.role.asc(), User.full_name.asc()).all()
    else:
        groups = current_user.teaching_groups
        group_ids = [group.id for group in groups]

        teachers = [current_user]

        if group_ids:
            all_users = User.query.filter(
                (User.role == 'teacher') & (User.id == current_user.id) |
                ((User.role == 'student') & (User.group_id.in_(group_ids)))
            ).order_by(User.role.asc(), User.full_name.asc()).all()
        else:
            all_users = [current_user]

        ungrouped_students = []

    return render_template(
        'admin/users.html',
        teachers=teachers,
        ungrouped_students=ungrouped_students,
        groups=groups,
        all_users=all_users,
    )

@admin_bp.route('/groups/<int:group_id>/students')
@login_required
def group_students(group_id):
    staff_required()

    group = StudentGroup.query.get_or_404(group_id)

    if not current_user.is_admin and group not in current_user.teaching_groups:
        abort(403)

    students = User.query.filter_by(
        role='student',
        group_id=group.id
    ).order_by(User.username.asc()).all()

    teachers = User.query.filter_by(role='teacher').order_by(User.full_name.asc()).all()

    return render_template(
        'admin/group_students.html',
        group=group,
        students=students,
        teachers=teachers,
    )

@admin_bp.route('/tests')
@login_required
def tests():
    staff_required()
    if current_user.role == 'admin':
        tests = Test.query.all()
    else:
        tests = Test.query.filter_by(teacher_id=current_user.id).all()
    csrf_form = FlaskForm()
    return render_template('admin/tests.html', tests=tests, question_type_labels=QUESTION_TYPE_LABEL_KEYS, csrf_form=csrf_form)


@admin_bp.route('/tests/create', methods=['GET', 'POST'])
@login_required
def create_test():
    staff_required()
    form = CreateTestForm()

    if form.validate_on_submit():
        test = Test(
            title=form.title.data.strip(),
            description=form.description.data.strip(),
            time_limit_minutes=form.time_limit_minutes.data,
            max_attempts=form.max_attempts.data,
            created_by=current_user.id,
            teacher_id=current_user.id
        )

        db.session.add(test)
        db.session.commit()

        log_action(
            action_type='create_test',
            entity_type='test',
            entity_id=test.id,
            description=f'Создал тест "{test.title}"',
        )
        db.session.commit()

        flash(_t('admin.test_created'), 'success')
        return redirect(url_for('admin.view_test', test_id=test.id))

    return render_template('admin/create_test.html', form=form)


@admin_bp.route('/tests/<int:test_id>')
@login_required
def view_test(test_id):
    staff_required()
    test = Test.query.get_or_404(test_id)

    if current_user.role != 'admin' and test.teacher_id != current_user.id:
        abort(403)

    if current_user.is_admin:
        students = User.query.filter_by(role='student').order_by(User.full_name.asc()).all()
        groups = StudentGroup.query.all()
    else:
        groups = current_user.teaching_groups
        students = []
    csrf_form = FlaskForm()

    if current_user.is_admin:
        visible_assignments = list(test.assignments)
    else:
        allowed_student_ids = set(get_teacher_student_ids())
        visible_assignments = [assignment for assignment in test.assignments if assignment.user_id in allowed_student_ids]

    assigned_user_ids = {assignment.user_id for assignment in visible_assignments}

    assignment_settings = {
        assignment.user_id: {
            'question_limit': assignment.question_limit or test.questions_per_student,
            'target_score': assignment.target_score or test.target_score,
        }
        for assignment in visible_assignments
    }

    return render_template(
        'admin/test_detail.html',
        test=test,
        students=students,
        groups=groups,
        assigned_user_ids=assigned_user_ids,
        assignment_settings=assignment_settings,
        question_type_labels=QUESTION_TYPE_LABEL_KEYS,
        csrf_form=csrf_form,
    )

@admin_bp.route('/tests/<int:test_id>/delete', methods=['POST'])
@login_required
def delete_test(test_id: int):
    staff_required()

    csrf_form = FlaskForm()
    if not csrf_form.validate_on_submit():
        abort(400)

    test = Test.query.get_or_404(test_id)

    if current_user.role != 'admin' and test.teacher_id != current_user.id:
        abort(403)

    title = test.title
    deleted_test_id = test.id

    db.session.delete(test)
    db.session.commit()

    log_action(
        action_type='delete_test',
        entity_type='test',
        entity_id=deleted_test_id,
        description=f'Удалил тест "{title}"',
    )
    db.session.commit()

    flash(_t('admin.test_deleted', title=title), 'info')
    return redirect(url_for('admin.tests'))

@admin_bp.route('/tests/<int:test_id>/duplicate', methods=['POST'])
@login_required
def duplicate_test(test_id: int):
    staff_required()

    original_test = Test.query.get_or_404(test_id)

    if current_user.role != 'admin' and original_test.teacher_id != current_user.id:
        abort(403)

    new_test = Test(
        title=f"{original_test.title} (копия)",
        description=original_test.description,
        time_limit_minutes=original_test.time_limit_minutes,
        max_attempts=original_test.max_attempts,
        questions_per_student=original_test.questions_per_student,
        target_score=original_test.target_score,
        created_by=current_user.id,
        teacher_id=current_user.id,
    )
    db.session.add(new_test)
    db.session.flush()

    for question in original_test.questions:
        new_question = Question(
            test_id=new_test.id,
            text=question.text,
            question_type=question.question_type,
            points=question.points,
        )
        db.session.add(new_question)
        db.session.flush()

        for option in question.answer_options:
            db.session.add(
                AnswerOption(
                    question_id=new_question.id,
                    text=option.text,
                    is_correct=option.is_correct,
                )
            )

        new_question.set_payload(question.payload or {})

    db.session.commit()

    log_action(
        action_type='duplicate_test',
        entity_type='test',
        entity_id=new_test.id,
        description=f'Создал копию теста "{original_test.title}" → "{new_test.title}"',
    )
    db.session.commit()

    flash('Тест продублирован', 'success')
    return redirect(url_for('admin.view_test', test_id=new_test.id))

@admin_bp.route('/tests/<int:test_id>/add_question', methods=['GET', 'POST'])
@login_required
def add_question(test_id):
    staff_required()
    test = Test.query.get_or_404(test_id)
    question_types = [choice[0] for choice in QUESTION_TYPE_CHOICES]

    if request.method == 'POST':
        question_data = parse_question_form_data(request.form)
        if 'error' in question_data:
            flash(_t(question_data['error']), 'warning')
            return render_template(
                'admin/add_question.html',
                test=test,
                question_types=question_types,
                data=request.form,
            )

        points = int(request.form.get('points', 1))

        question = Question(
            test_id=test.id,
            text='',
            question_type='single_choice',
            points=points,
        )
        db.session.add(question)
        db.session.flush()

        save_question_from_data(question, question_data, points)
        db.session.commit()

        log_action(
            action_type='add_question',
            entity_type='question',
            entity_id=question.id,
            description=f'Добавил вопрос в тест "{test.title}": "{question.text[:120]}"',
        )
        db.session.commit()

        flash(_t('admin.question_added'), 'success')
        return redirect(url_for('admin.view_test', test_id=test.id))

    return render_template(
        'admin/add_question.html',
        test=test,
        question_types=question_types,
        data=None,
    )
@admin_bp.route('/questions/<int:question_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_question(question_id):
    staff_required()

    question = Question.query.get_or_404(question_id)
    test = question.test

    if current_user.role != 'admin' and test.teacher_id != current_user.id:
        abort(403)

    question_types = [choice[0] for choice in QUESTION_TYPE_CHOICES]

    if request.method == 'POST':
        question_data = parse_question_form_data(request.form)
        if 'error' in question_data:
            flash(_t(question_data['error']), 'warning')
            return render_template(
                'admin/edit_question.html',
                test=test,
                question=question,
                question_types=question_types,
                data=request.form,
            )

        points = int(request.form.get('points', 1))
        save_question_from_data(question, question_data, points)
        db.session.commit()

        log_action(
            action_type='edit_question',
            entity_type='question',
            entity_id=question.id,
            description=f'Изменил вопрос в тесте "{test.title}": "{question.text[:120]}"',
        )
        db.session.commit()

        flash('Вопрос обновлён', 'success')
        return redirect(url_for('admin.view_test', test_id=test.id))

    form_data = {
        'text': question.text,
        'question_type': question.normalized_type,
        'points': question.points,
    }

    if question.normalized_type in {'single_choice', 'multiple_choice'}:
        correct_options = []
        for index, option in enumerate(question.answer_options, start=1):
            form_data[f'option_{index}'] = option.text
            if option.is_correct:
                correct_options.append(str(index))
        form_data['correct_options'] = correct_options

    if question.normalized_type == 'text':
        form_data['text_correct_answer'] = question.payload.get('correct_text', '')

    return render_template(
        'admin/edit_question.html',
        test=test,
        question=question,
        question_types=question_types,
        data=form_data,
    )


@admin_bp.route('/questions/<int:question_id>/duplicate', methods=['POST'])
@login_required
def duplicate_question(question_id):
    staff_required()

    question = Question.query.get_or_404(question_id)
    test = question.test

    if current_user.role != 'admin' and test.teacher_id != current_user.id:
        abort(403)

    new_question = Question(
        test_id=test.id,
        text=f"{question.text} (копия)",
        question_type=question.question_type,
        points=question.points,
    )
    db.session.add(new_question)
    db.session.flush()

    for option in question.answer_options:
        db.session.add(
            AnswerOption(
                question_id=new_question.id,
                text=option.text,
                is_correct=option.is_correct,
            )
        )

    new_question.set_payload(question.payload or {})
    db.session.commit()

    log_action(
        action_type='duplicate_question',
        entity_type='question',
        entity_id=new_question.id,
        description=f'Продублировал вопрос в тесте "{test.title}": "{question.text[:120]}"',
    )
    db.session.commit()

    flash('Вопрос продублирован', 'success')
    return redirect(url_for('admin.view_test', test_id=test.id))


@admin_bp.route('/tests/<int:test_id>/assign', methods=['POST'])
@login_required
def assign_test(test_id):
    staff_required()
    test = Test.query.get_or_404(test_id)

    if current_user.role != 'admin' and test.teacher_id != current_user.id:
        abort(403)

    selected_students = request.form.getlist('student_ids')
    selected_groups = request.form.getlist('group_ids')
    allowed_group_ids = set(get_teacher_group_ids()) if not current_user.is_admin else None
    allowed_student_ids = set(get_teacher_student_ids()) if not current_user.is_admin else None

    assignments_map: dict[int, dict[str, int]] = {}

    # Назначение целым группам
    for group_id in selected_groups:
        if not group_id.isdigit():
            continue

        group_id_int = int(group_id)

        if allowed_group_ids is not None and group_id_int not in allowed_group_ids:
            continue

        group = StudentGroup.query.get(group_id_int)
        if not group:
            continue

        question_limit_raw = request.form.get(f'group_question_limit_{group.id}', str(test.questions_per_student)).strip()
        target_score_raw = request.form.get(f'group_target_score_{group.id}', str(test.target_score)).strip()

        question_limit = int(question_limit_raw) if question_limit_raw.isdigit() else test.questions_per_student
        target_score = int(target_score_raw) if target_score_raw.isdigit() else test.target_score

        if question_limit <= 0:
            question_limit = test.questions_per_student
        if target_score <= 0:
            target_score = test.target_score

        for student in group.students:
            assignments_map[student.id] = {
                'question_limit': question_limit,
                'target_score': target_score,
            }

    # Назначение отдельным ученикам
    for student_id in selected_students:
        if not student_id.isdigit():
            continue

        student_id_int = int(student_id)
        if allowed_student_ids is not None and student_id_int not in allowed_student_ids:
            continue

        student = User.query.filter_by(id=student_id_int, role='student').first()
        if not student:
            continue

        # если ученик уже попал через группу — оставляем групповые настройки
        if student_id_int in assignments_map:
            continue

        question_limit_raw = request.form.get(
            f'student_question_limit_{student_id_int}',
            str(test.questions_per_student)
        ).strip()
        target_score_raw = request.form.get(
            f'student_target_score_{student_id_int}',
            str(test.target_score)
        ).strip()

        question_limit = int(question_limit_raw) if question_limit_raw.isdigit() else test.questions_per_student
        target_score = int(target_score_raw) if target_score_raw.isdigit() else test.target_score

        if question_limit <= 0:
            question_limit = test.questions_per_student
        if target_score <= 0:
            target_score = test.target_score

        assignments_map[student_id_int] = {
            'question_limit': question_limit,
            'target_score': target_score,
        }

    TestAssignment.query.filter_by(test_id=test.id).delete()

    for user_id, settings in assignments_map.items():
        db.session.add(
            TestAssignment(
                user_id=user_id,
                test_id=test.id,
                question_limit=settings['question_limit'],
                target_score=settings['target_score'],
            )
        )

    db.session.commit()

    log_action(
        action_type='assign_test',
        entity_type='test',
        entity_id=test.id,
        description=f'Сохранил назначения для теста "{test.title}". Назначений: {len(assignments_map)}',
    )
    db.session.commit()

    flash(_t('admin.assignments_saved'), 'success')
    return redirect(url_for('admin.view_test', test_id=test.id))

@admin_bp.route('/logs')
@login_required
def action_logs():
    staff_required()

    if current_user.is_admin:
        logs = ActionLog.query.order_by(ActionLog.created_at.desc()).limit(300).all()
    else:
        logs = ActionLog.query.filter_by(actor_user_id=current_user.id).order_by(ActionLog.created_at.desc()).limit(300).all()

    return render_template('admin/action_logs.html', logs=logs)

@admin_bp.route('/results')
@login_required
def results():
    staff_required()

    attempts_query = Attempt.query.filter(Attempt.finished_at.isnot(None))
    student_id = request.args.get('student_id', type=int)
    test_id = request.args.get('test_id', type=int)
    group_id = request.args.get('group_id', type=int)

    if current_user.is_admin:
        students = User.query.filter_by(role='student').order_by(User.full_name.asc()).all()
        tests_list = Test.query.order_by(Test.title.asc()).all()
    else:
        teacher_student_ids = get_teacher_student_ids()
        teacher_test_ids = get_teacher_test_ids()

        if teacher_test_ids:
            attempts_query = attempts_query.filter(Attempt.test_id.in_(teacher_test_ids))
            tests_list = Test.query.filter(Test.id.in_(teacher_test_ids)).order_by(Test.title.asc()).all()
        else:
            attempts_query = attempts_query.filter(False)
            tests_list = []

        if teacher_student_ids:
            students = User.query.filter(
                User.role == 'student',
                User.id.in_(teacher_student_ids)
            ).order_by(User.full_name.asc()).all()
        else:
            students = []
    if student_id:
        attempts_query = attempts_query.filter(Attempt.user_id == student_id)

    if test_id:
        attempts_query = attempts_query.filter(Attempt.test_id == test_id)

    if group_id:
        attempts_query = attempts_query.join(User).filter(User.group_id == group_id)

    # ВСЕГДА В КОНЦЕ
    attempts = attempts_query.order_by(Attempt.finished_at.desc()).all()

    for attempt in attempts:
        attempt.risk_reasons = build_attempt_risk_reasons(attempt)

    if request.args.get('export') == 'xlsx':
        return export_attempts_to_excel(attempts)

    groups = StudentGroup.query.order_by(StudentGroup.name.asc()).all()

    return render_template(
        'admin/results.html',
        attempts=attempts,
        students=students,
        tests=tests_list,
        groups=groups,
        selected_student_id=student_id,
        selected_test_id=test_id,
        selected_group_id=group_id
    )


@admin_bp.route('/seed', methods=['POST'])
@login_required
def seed():
    admin_required()
    seed_demo_data()
    flash(_t('admin.seed_success'), 'success')
    return redirect(url_for('admin.admin_dashboard'))



def export_attempts_to_excel(attempts):
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = 'Results'
    sheet.append([
        "Студент",
        "Группа",
        "Тест",
        "Начало",
        "Завершение",
        "Баллы",
        "Переключения вкладок",
    ])
    for attempt in attempts:
        sheet.append([
            attempt.user.full_name,
            attempt.user.group.name if attempt.user.group else "-",
            attempt.test.title,
            attempt.started_at.strftime('%Y-%m-%d %H:%M'),
            attempt.finished_at.strftime('%Y-%m-%d %H:%M') if attempt.finished_at else '',
            f'{attempt.earned_score} / {attempt.max_score}',
            attempt.tab_switch_count or 0,
        ])

    output = BytesIO()
    workbook.save(output)
    output.seek(0)
    return send_file(
        output,
        as_attachment=True,
        download_name='results.xlsx',
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )



def create_tests_from_parsed_data(parsed_tests, time_limit_minutes: int, max_attempts: int, created_by: int, teacher_id: int) -> tuple[int, list[int]]:
    created_count = 0
    created_ids: list[int] = []

    for parsed_test in parsed_tests:
        test = Test(
            title=parsed_test.title,
            description=parsed_test.description,
            time_limit_minutes=time_limit_minutes,
            max_attempts=max_attempts,
            created_by=created_by,
            teacher_id=teacher_id,
        )

        db.session.add(test)
        db.session.flush()
        created_ids.append(int(test.id))

        for parsed_question in parsed_test.questions:
            question = Question(
                test=test,
                text=parsed_question.text,
                question_type=parsed_question.question_type,
                points=parsed_question.points,
                image_path=parsed_question.image_path,
            )
            db.session.add(question)
            db.session.flush()

            if parsed_question.answers:
                for answer_text, is_correct in parsed_question.answers:
                    db.session.add(
                        AnswerOption(
                            question=question,
                            text=answer_text,
                            is_correct=is_correct,
                        )
                    )

            if parsed_question.payload:
                question.set_payload(parsed_question.payload)

        created_count += 1

    db.session.commit()
    return created_count, created_ids


def parse_question_form_data(form) -> dict:
    question_text = form.get('text', '').strip()
    question_type = form.get('question_type', 'single_choice').strip()

    if not question_text:
        return {'error': 'admin.question_text_required'}

    if question_type in {'single_choice', 'multiple_choice'}:
        answer_options = []
        correct_indexes = set(form.getlist('correct_options'))
        for index in range(1, 7):
            option_text = form.get(f'option_{index}', '').strip()
            if option_text:
                answer_options.append((option_text, str(index) in correct_indexes))

        if len(answer_options) < 2:
            return {'error': 'admin.need_two_options'}
        if not any(is_correct for _, is_correct in answer_options):
            return {'error': 'admin.need_one_correct'}
        if question_type == 'single_choice' and sum(1 for _, is_correct in answer_options if is_correct) != 1:
            return {'error': 'admin.single_only_one_correct'}

        return {'text': question_text, 'question_type': question_type, 'answer_options': answer_options, 'payload': {}}

    if question_type == 'text':
        correct_text = form.get('text_correct_answer', '').strip()
        if not correct_text:
            return {'error': 'admin.text_answer_required'}
        return {
            'text': question_text,
            'question_type': question_type,
            'answer_options': [],
            'payload': {'correct_text': correct_text},
        }

    if question_type == 'matching':
        left_items = [value.strip() for value in form.getlist('matching_left[]') if value.strip()]
        right_items = [value.strip() for value in form.getlist('matching_right[]') if value.strip()]
        if len(left_items) < 2 or len(left_items) != len(right_items):
            return {'error': 'admin.matching_pairs_required'}
        pairs = [{'left': left, 'right': right} for left, right in zip(left_items, right_items)]
        return {'text': question_text, 'question_type': question_type, 'answer_options': [], 'payload': {'pairs': pairs}}

    if question_type == 'ordering':
        items = [value.strip() for value in form.getlist('ordering_items[]') if value.strip()]
        if len(items) < 2:
            return {'error': 'admin.ordering_items_required'}
        return {'text': question_text, 'question_type': question_type, 'answer_options': [], 'payload': {'items': items}}

    return {'error': 'admin.unsupported_question_type'}

def save_question_from_data(question: Question, question_data: dict, points: int) -> None:
    question.text = question_data['text']
    question.question_type = question_data['question_type']
    question.points = points

    AnswerOption.query.filter_by(question_id=question.id).delete()

    for option_text, is_correct in question_data['answer_options']:
        db.session.add(
            AnswerOption(
                question_id=question.id,
                text=option_text,
                is_correct=is_correct,
            )
        )

    question.set_payload(question_data['payload'] or {})

def save_question_image(file) -> str | None:
    if not file or not getattr(file, 'filename', ''):
        return None

    filename = secure_filename(file.filename)
    if not filename:
        return None

    unique_name = f"{uuid4().hex}_{filename}"
    upload_dir = Path(current_app.static_folder) / 'uploads' / 'questions'
    upload_dir.mkdir(parents=True, exist_ok=True)

    save_path = upload_dir / unique_name
    file.save(save_path)

    return f"uploads/questions/{unique_name}"


def save_question_from_data(question: Question, question_data: dict, points: int) -> None:
    question.text = question_data['text']
    question.question_type = question_data['question_type']
    question.points = points

    image_file = request.files.get('image')
    if image_file and getattr(image_file, 'filename', ''):
        question.image_path = save_question_image(image_file)

    AnswerOption.query.filter_by(question_id=question.id).delete()

    for option_text, is_correct in question_data['answer_options']:
        db.session.add(
            AnswerOption(
                question_id=question.id,
                text=option_text,
                is_correct=is_correct,
            )
        )

    question.set_payload(question_data['payload'] or {})

@admin_bp.route('/users/create', methods=['GET', 'POST'])
@login_required
def create_user():
    admin_required()
    form = CreateUserForm()

    groups = StudentGroup.query.order_by(StudentGroup.name.asc()).all()
    form.group_id.choices = [(0, 'Без группы')] + [(group.id, group.name) for group in groups]

    if form.validate_on_submit():
        username = form.username.data.strip()
        if User.query.filter_by(username=username).first():
            flash('Пользователь с таким логином уже существует', 'danger')
            return render_template('admin/create_user.html', form=form)

        selected_group_id = form.group_id.data if form.role.data == 'student' and form.group_id.data != 0 else None

        user = User(
            username=username,
            full_name=form.full_name.data.strip(),
            role=form.role.data,
            group_id=selected_group_id,
        )
        user.set_password(form.password.data)

        db.session.add(user)
        db.session.commit()

        flash('Пользователь создан', 'success')
        return redirect(url_for('admin.users'))

    return render_template('admin/create_user.html', form=form)
def admin_required():
    if not current_user.is_authenticated or not current_user.is_admin:
        abort(403)

@admin_bp.route('/users/<int:user_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_user(user_id):
    admin_required()

    user = User.query.get_or_404(user_id)
    form = EditUserForm(obj=user)

    groups = StudentGroup.query.order_by(StudentGroup.name.asc()).all()
    form.group_id.choices = [(0, 'Без группы')] + [(group.id, group.name) for group in groups]

    if request.method == 'GET':
        form.group_id.data = user.group_id or 0

    if form.validate_on_submit():
        username = form.username.data.strip()
        existing_user = User.query.filter(User.username == username, User.id != user.id).first()
        if existing_user:
            flash('Пользователь с таким логином уже существует', 'danger')
            return render_template('admin/edit_user.html', form=form, user=user)

        user.full_name = form.full_name.data.strip()
        user.username = username
        user.role = form.role.data

        if user.role == 'student':
            user.group_id = form.group_id.data if form.group_id.data != 0 else None
        else:
            user.group_id = None

        new_password = (form.password.data or '').strip()
        if new_password:
            user.set_password(new_password)

        db.session.commit()
        flash('Данные пользователя обновлены', 'success')
        return redirect(url_for('admin.users'))

    return render_template('admin/edit_user.html', form=form, user=user)

@admin_bp.route('/groups/upload', methods=['GET', 'POST'])
@login_required
def upload_group():
    admin_required()
    form = UploadGroupForm()

    if form.validate_on_submit():
        file = form.file.data
        original_name = secure_filename(file.filename or 'group.xlsx')
        extension = Path(original_name).suffix.lower()
        unique_name = f'{uuid4().hex}{extension}'
        upload_dir = Path(current_app.instance_path) / current_app.config['UPLOAD_FOLDER']
        upload_dir.mkdir(parents=True, exist_ok=True)
        saved_path = upload_dir / unique_name
        file.save(saved_path)

        try:
            group_name, students_data = parse_group_excel(str(saved_path))
        except ParseError as error:
            saved_path.unlink(missing_ok=True)
            flash(str(error), 'danger')
            return render_template('admin/upload_group.html', form=form)

        group = StudentGroup.query.filter_by(name=group_name).first()
        if not group:
            group = StudentGroup(name=group_name)
            db.session.add(group)
            db.session.flush()

        created_count = 0

        for item in students_data:
            existing_user = User.query.filter_by(username=item['username']).first()
            if existing_user:
                continue

            user = User(
                username=item['username'],
                full_name=item['full_name'],
                role='student',
                group=group,
            )
            user.set_password(item['password'])
            db.session.add(user)
            created_count += 1

        db.session.commit()
        saved_path.unlink(missing_ok=True)

        flash(f'Группа {group_name} импортирована. Добавлено студентов: {created_count}', 'success')
        return redirect(url_for('admin.groups'))

    return render_template('admin/upload_group.html', form=form)

@admin_bp.route('/groups')
@login_required
def groups():
    staff_required()

    if current_user.is_admin:
        groups_list = StudentGroup.query.order_by(StudentGroup.name.asc()).all()
    else:
        groups_list = current_user.teaching_groups

    teachers = User.query.filter_by(role='teacher').order_by(User.full_name.asc()).all()
    return render_template('admin/groups.html', groups=groups_list, teachers=teachers)

@admin_bp.route('/groups/<int:group_id>/assign', methods=['POST'])
@login_required
def assign_group(group_id):
    admin_required()

    group = StudentGroup.query.get_or_404(group_id)
    teacher_ids = request.form.getlist('teacher_ids')

    group.teachers.clear()

    for teacher_id in teacher_ids:
        teacher = User.query.filter_by(id=int(teacher_id), role='teacher').first()
        if teacher:
            group.teachers.append(teacher)

    db.session.commit()
    flash(f'Назначения для группы {group.name} сохранены', 'success')
    return redirect(url_for('admin.groups'))
@admin_bp.route('/users/<int:user_id>/delete', methods=['POST'])
@login_required
def delete_user(user_id):
    admin_required()

    user = User.query.get_or_404(user_id)

    if user.id == current_user.id:
        flash('Нельзя удалить самого себя', 'warning')
        return redirect(url_for('admin.users'))

    db.session.delete(user)
    db.session.commit()

    flash('Пользователь удален', 'success')
    return redirect(url_for('admin.users'))

from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _add_answer_count_line(paragraph, question) -> None:
    normalized = getattr(question, 'normalized_type', question.question_type)

    if normalized == 'single_choice':
        text = "Toʻgʻri javoblar soni: 1"
    elif normalized == 'multiple_choice':
        correct_count = sum(1 for option in question.answer_options if option.is_correct)
        text = f"Toʻgʻri javoblar soni: {correct_count}"
    else:
        text = "Toʻgʻri javoblar soni: 1"

    run = paragraph.add_run(text + "\n")
    run.italic = True
    run.font.size = Pt(10)



def _render_choice_question(paragraph, question) -> None:
    options = list(question.answer_options)
    random.shuffle(options)

    for opt in options:
        square_run = paragraph.add_run("□ ")
        square_run.font.size = Pt(13)

        text_run = paragraph.add_run(f"{opt.text}\n")
        text_run.font.size = Pt(12)


def _render_text_question(paragraph, question) -> None:
    paragraph.add_run("Javob: ").font.size = Pt(12)
    line_run = paragraph.add_run("_______________________________\n")
    line_run.font.size = Pt(12)

    paragraph.add_run("\n").font.size = Pt(6)

    paragraph.add_run("_______________________________________________\n").font.size = Pt(12)
    paragraph.add_run("_______________________________________________\n").font.size = Pt(12)


def _render_matching_question(paragraph, question) -> None:
    pairs = question.payload.get('pairs', []) if question.payload else []

    if not pairs:
        paragraph.add_run("Juftliklar topilmadi.\n").font.size = Pt(11)
        return

    left_items = []
    right_items = []

    for pair in pairs:
        left = pair.left if hasattr(pair, 'left') else pair.get('left', '')
        right = pair.right if hasattr(pair, 'right') else pair.get('right', '')
        left_items.append(str(left))
        right_items.append(str(right))

    shuffled_right = list(right_items)
    random.shuffle(shuffled_right)

    paragraph.add_run("Chap ustun:\n").bold = True
    for idx, item in enumerate(left_items, start=1):
        run = paragraph.add_run(f"  {idx}) {item}\n")
        run.font.size = Pt(12)

    paragraph.add_run("\nOʻng ustun:\n").bold = True
    for idx, item in enumerate(shuffled_right, start=1):
        label = chr(64 + idx) if idx <= 26 else str(idx)
        run = paragraph.add_run(f"  {label}) {item}\n")
        run.font.size = Pt(12)

    paragraph.add_run("\nJavob: ").font.size = Pt(11)
    paragraph.add_run("1-__  2-__  3-__  4-__\n").font.size = Pt(11)


def _render_ordering_question(paragraph, question) -> None:
    items = question.payload.get('items', []) if question.payload else []

    if not items:
        paragraph.add_run("Elementlar topilmadi.\n").font.size = Pt(11)
        return

    shuffled_items = list(items)
    random.shuffle(shuffled_items)

    paragraph.add_run("Tartiblang:\n").bold = True
    for idx, item in enumerate(shuffled_items, start=1):
        run = paragraph.add_run(f"  □ {item}\n")
        run.font.size = Pt(12)

    paragraph.add_run("\nJavob tartibi: __________________________\n").font.size = Pt(11)


def _render_question_content(paragraph, question) -> None:
    normalized = getattr(question, 'normalized_type', question.question_type)

    _add_answer_count_line(paragraph, question)

    if normalized in ('single_choice', 'multiple_choice'):
        _render_choice_question(paragraph, question)
    elif normalized == 'text':
        _render_text_question(paragraph, question)
    elif normalized == 'matching':
        _render_matching_question(paragraph, question)
    elif normalized == 'ordering':
        _render_ordering_question(paragraph, question)
    else:
        paragraph.add_run("Bu savol turi uchun eksport shakli topilmadi.\n").font.size = Pt(11)

@admin_bp.route('/tests/<int:test_id>/generate_variants', methods=['POST'])
@login_required
def generate_variants(test_id):
    staff_required()

    test = Test.query.get_or_404(test_id)

    variants_count = int(request.form.get('variants_count', 5))
    questions_per_variant = int(request.form.get('questions_per_variant', 20))
    target_score = int(request.form.get('target_score', 20))

    questions = list(test.questions)

    doc = Document()

    # Небольшая настройка полей
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    for i in range(variants_count):
        title = doc.add_paragraph(f"ВАРИАНТ {i + 1}")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(14)

        doc.add_paragraph(
            "________________________________________________________________________________(Oʻqish joyi) "
            "_____________________________________________________________________ fakulteti, "
            "_____________________________________ yoʻnalishi,_________________________________guruhi, "
            "talabasi(FIO) ________________________________________________________________ "
        )
        doc.add_paragraph("")

        selected = choose_questions_exact(
            questions=questions,
            question_limit=questions_per_variant,
            target_score=target_score,
            seed=f"{i}-{random.random()}",
        )

        if not selected:
            continue

        table = doc.add_table(rows=0, cols=2)
        table.autofit = True

        for idx in range(0, len(selected), 2):
            row = table.add_row().cells

            # Левая колонка
            q1 = selected[idx]
            row[0].text = ""
            p1 = row[0].paragraphs[0]

            q1_run = p1.add_run(f"{idx + 1}. {q1.text}\n")
            q1_run.font.size = Pt(11)
            q1_run.bold = False

            _render_question_content(p1, q1)

            # Правая колонка
            if idx + 1 < len(selected):
                q2 = selected[idx + 1]
                row[1].text = ""
                p2 = row[1].paragraphs[0]

                q2_run = p2.add_run(f"{idx + 2}. {q2.text}\n")
                q2_run.font.size = Pt(11)
                q2_run.bold = False

                _render_question_content(p2, q2)
            else:
                row[1].text = ""

        if i < variants_count - 1:
            doc.add_page_break()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"variants_test_{test.id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

def save_question_image(file):
    if not file:
        return None

    filename = secure_filename(file.filename)
    unique_name = f"{uuid4().hex}_{filename}"

    upload_path = Path("app/static/uploads")
    upload_path.mkdir(parents=True, exist_ok=True)

    file_path = upload_path / unique_name
    file.save(file_path)

    return f"uploads/{unique_name}"
@admin_bp.route('/tests/<int:test_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_test(test_id):
    staff_required()

    test = Test.query.get_or_404(test_id)

    if current_user.role != 'admin' and test.teacher_id != current_user.id:
        abort(403)

    form = CreateTestForm(obj=test)

    if form.validate_on_submit():
        test.title = form.title.data
        test.description = form.description.data
        test.time_limit_minutes = form.time_limit_minutes.data
        test.max_attempts = form.max_attempts.data
        test.target_score = form.target_score.data
        test.questions_per_student = form.questions_per_student.data

        db.session.commit()

        flash('Test yangilandi', 'success')
        return redirect(url_for('admin.view_test', test_id=test.id))

    return render_template('admin/edit_test.html', form=form, test=test)